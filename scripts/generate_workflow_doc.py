"""Generate TMS Pro Workflow Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUTPUT = r"C:\Users\PLC\Desktop\TMS\docs\TMS_Pro_Workflow.docx"

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")

def main():
    doc = Document()

    # Title
    title = doc.add_heading("TMS Pro — System Workflow Document", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Transport Management System | Version 1.0 | {date.today().strftime('%d %B %Y')}")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # 1. Introduction
    add_heading(doc, "1. Introduction", 1)
    add_para(doc,
        "TMS Pro is an enterprise Transport Management System designed for logistics and transport "
        "companies. It manages the complete lifecycle from customer booking through lorry receipt (LR) "
        "generation, fleet operations, expense tracking, accounting, and business reporting."
    )
    add_para(doc,
        "This document describes the system architecture, user workflows, business process flows, "
        "installation workflow, and data flow between modules."
    )

    # 2. System Architecture
    add_heading(doc, "2. System Architecture", 1)
    add_para(doc, "TMS Pro follows a three-tier architecture:", bold=True)
    add_bullet(doc, "Presentation Layer: React 19 + Vite + Tailwind CSS (Port 5173)")
    add_bullet(doc, "Application Layer: ASP.NET Core 8 REST API with JWT Authentication (Port 5000)")
    add_bullet(doc, "Data Layer: PostgreSQL Database (Port 5432)")
    doc.add_paragraph()
    add_para(doc, "Request Flow:", bold=True)
    add_numbered(doc, "User interacts with React frontend in the browser.")
    add_numbered(doc, "Frontend sends HTTP requests to /api/* (proxied to backend).")
    add_numbered(doc, "API validates JWT token and processes business logic.")
    add_numbered(doc, "Entity Framework Core reads/writes data to PostgreSQL.")
    add_numbered(doc, "JSON response returned to frontend and displayed in UI.")

    # 3. Authentication Workflow
    add_heading(doc, "3. Authentication Workflow", 1)
    add_numbered(doc, "User opens application URL (http://localhost:5173).")
    add_numbered(doc, "If not logged in, user is redirected to /login page.")
    add_numbered(doc, "User enters username and password (default: admin / admin123).")
    add_numbered(doc, "Frontend sends POST /api/auth/login to backend.")
    add_numbered(doc, "Backend validates credentials against users table (BCrypt hash).")
    add_numbered(doc, "On success, JWT token is returned and stored in localStorage.")
    add_numbered(doc, "User is redirected to Dashboard; all API calls include Bearer token.")
    add_numbered(doc, "On token expiry (401), user is logged out and returned to login.")

    # 4. Master Data Setup Workflow
    add_heading(doc, "4. Master Data Setup Workflow (Initial Configuration)", 1)
    add_para(doc, "Before daily operations, configure master records in this recommended order:")
    add_numbered(doc, "Settings → Company Profile: Enter company name, GSTIN, PAN, address, financial year.")
    add_numbered(doc, "Customers → Add New: Register consignors/consignees with credit limits.")
    add_numbered(doc, "Vendors → Add New: Register fuel, maintenance, and other suppliers.")
    add_numbered(doc, "Vehicles → Add New: Register fleet with registration, insurance, fitness dates.")
    add_numbered(doc, "Drivers → Add New: Register drivers with license, salary, and contact details.")
    add_numbered(doc, "Accounting → Ledger Master: Create ledger accounts (Cash, Bank, Income, Expense).")

    # 5. Core Business Workflow
    add_heading(doc, "5. Core Business Workflow — Booking to Trip Completion", 1)
    add_para(doc, "This is the primary operational workflow for transport operations:", bold=True)

    add_heading(doc, "5.1 Create Booking", 2)
    add_numbered(doc, "Navigate to Bookings → Add New Booking.")
    add_numbered(doc, "Select customer, enter route (From/To), material, quantity.")
    add_numbered(doc, "Assign vehicle and driver from active fleet.")
    add_numbered(doc, "Enter freight amount, advance, payment status (Paid/Partial/Unpaid).")
    add_numbered(doc, "Save → API creates record in bookings table (Status: Pending/Confirmed).")

    add_heading(doc, "5.2 Generate Lorry Receipt (LR)", 2)
    add_numbered(doc, "Navigate to LR → Generate New LR.")
    add_numbered(doc, "Enter consignor, consignee, route, vehicle, driver details.")
    add_numbered(doc, "Enter freight, GST, hamali, loading/unloading charges, advance.")
    add_numbered(doc, "System calculates balance = total charges − advance.")
    add_numbered(doc, "Save → API creates LR in lorry_receipts table; LR number auto-generated.")
    add_numbered(doc, "LR links to booking (optional) and updates sales register data.")

    add_heading(doc, "5.3 Trip Execution & Status Updates", 2)
    add_numbered(doc, "Booking status progresses: Pending → Confirmed → In Transit → Delivered.")
    add_numbered(doc, "Vehicle status may change: Active → On Trip → Active.")
    add_numbered(doc, "Dashboard shows trip analysis, recent trips, and pending LR alerts.")

    add_heading(doc, "5.4 Record Expenses", 2)
    add_numbered(doc, "Navigate to Expenses → Add New Expense.")
    add_numbered(doc, "Select category: Fuel, Salary, Toll, Maintenance, Office, etc.")
    add_numbered(doc, "Link to vehicle and/or vendor; enter amount and payment mode.")
    add_numbered(doc, "Save → Expense recorded; feeds expense reports and P&L.")

    add_heading(doc, "5.5 Payment Collection", 2)
    add_numbered(doc, "When customer pays freight: Accounting → Voucher Entry → Receipt Voucher.")
    add_numbered(doc, "Debit: Cash/Bank account | Credit: Customer/Freight Income ledger.")
    add_numbered(doc, "Update booking payment status to Paid or Partial.")
    add_numbered(doc, "Receipt appears in Receipt Register and Cash/Bank Book.")

    # 6. Accounting Workflow
    add_heading(doc, "6. Accounting Workflow", 1)
    add_heading(doc, "6.1 Voucher Entry", 2)
    add_bullet(doc, "Payment Voucher: Pay vendors, fuel, salary (Debit expense, Credit cash/bank).")
    add_bullet(doc, "Receipt Voucher: Collect freight from customers (Debit cash/bank, Credit income).")
    add_bullet(doc, "Journal Voucher: Adjustments, GST entries, transfers between accounts.")
    add_bullet(doc, "Contra Voucher: Cash to bank or bank to cash transfers.")

    add_heading(doc, "6.2 Registers & Books", 2)
    add_bullet(doc, "Day Book: All vouchers for the day in chronological order.")
    add_bullet(doc, "Cash Book: Cash receipts and payments with running balance.")
    add_bullet(doc, "Bank Book: Bank deposits and withdrawals with running balance.")
    add_bullet(doc, "Sales Register: All LR/freight sales with GST.")
    add_bullet(doc, "Purchase Register: All vendor/expense purchases with GST.")

    add_heading(doc, "6.3 Financial Statements", 2)
    add_numbered(doc, "Trial Balance: Verify total debits equal total credits.")
    add_numbered(doc, "Profit & Loss: Compare freight income vs operating expenses.")
    add_numbered(doc, "Balance Sheet: Assets, liabilities, and capital position.")
    add_numbered(doc, "Outstanding Report: Customer receivables and vendor payables aging.")
    add_numbered(doc, "GST Reports: Input GST vs output GST, net payable.")

    # 7. Reporting Workflow
    add_heading(doc, "7. Reporting & Dashboard Workflow", 1)
    add_heading(doc, "7.1 Dashboard", 2)
    add_bullet(doc, "Overview tab: KPI cards (vehicles, drivers, income, expenses, cash/bank balance).")
    add_bullet(doc, "Analytics tab: 14 chart widgets (revenue, expenses, fleet, routes, drivers).")
    add_bullet(doc, "Recent Bookings and Recent Trips tabs with live API data.")
    add_bullet(doc, "Alerts panel: Unpaid bookings, maintenance vehicles, high receivables.")

    add_heading(doc, "7.2 Reports Hub", 2)
    add_bullet(doc, "Trip Report: LR-wise freight, expense, profit analysis.")
    add_bullet(doc, "Vehicle/Driver/Customer/Vendor Reports: Entity performance summaries.")
    add_bullet(doc, "Income & Expense Reports: Monthly breakdown by category.")
    add_bullet(doc, "Cash Flow Report: Monthly inflow, outflow, and net cash position.")

    # 8. Data Flow Diagram (text)
    add_heading(doc, "8. End-to-End Data Flow", 1)
    flow = doc.add_paragraph()
    flow.add_run(
        "Customer Booking  →  LR Generation  →  Trip Execution  →  Expense Recording\n"
        "       ↓                    ↓                  ↓                    ↓\n"
        "  bookings table    lorry_receipts      vehicles/drivers     expenses table\n"
        "       ↓                    ↓                  ↓                    ↓\n"
        "  Sales Register    Dashboard Stats     Fleet Reports        P&L / Expense Report\n"
        "       ↓                    ↓                  ↓                    ↓\n"
        "  Receipt Voucher  →  Cash/Bank Book  →  Outstanding  →  Balance Sheet / Trial Balance"
    ).font.name = "Consolas"

    # 9. Installation Workflow
    add_heading(doc, "9. Installation & Deployment Workflow", 1)
    add_heading(doc, "9.1 Prerequisites", 2)
    add_bullet(doc, "Node.js 18+ and npm")
    add_bullet(doc, ".NET 8 SDK")
    add_bullet(doc, "PostgreSQL 14+ with psql command-line tools")
    add_bullet(doc, "pgAdmin 4 (optional, for database management)")

    add_heading(doc, "9.2 Database Setup", 2)
    add_numbered(doc, "Install PostgreSQL and note postgres user password.")
    add_numbered(doc, "Create database: CREATE DATABASE tms_pro;")
    add_numbered(doc, "Run database/schema.sql to create tables.")
    add_numbered(doc, "Run database/seed.sql to seed master data and admin user.")
    add_numbered(doc, "Run database/seed_accounting.sql for ledger accounts and vouchers.")

    add_heading(doc, "9.3 Backend Setup", 2)
    add_numbered(doc, "Update DefaultConnection in backend/Tms.Api/appsettings.json.")
    add_numbered(doc, "cd backend/Tms.Api && dotnet restore && dotnet run.")
    add_numbered(doc, "Verify API at http://localhost:5000/swagger.")
    add_numbered(doc, "DbSeeder creates admin user on first run if not present.")

    add_heading(doc, "9.4 Frontend Setup", 2)
    add_numbered(doc, "npm install in project root.")
    add_numbered(doc, "Configure VITE_API_URL=/api in .env file.")
    add_numbered(doc, "npm run dev — opens http://localhost:5173.")
    add_numbered(doc, "Login with admin / admin123.")

    # 10. User Roles
    add_heading(doc, "10. User Roles & Access", 1)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Role", "Access Level", "Primary Modules"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    roles = [
        ("Super Admin", "Full access", "All modules + Settings"),
        ("Accountant", "Accounting focus", "Accounting, Reports, Vouchers"),
        ("Operations Manager", "Operations focus", "Bookings, LR, Fleet, Expenses"),
        ("Driver Coordinator", "Fleet focus", "Drivers, Vehicles, LR"),
        ("Viewer", "Read-only", "Dashboard, Reports (view only)"),
    ]
    for i, (role, access, modules) in enumerate(roles, 1):
        table.rows[i].cells[0].text = role
        table.rows[i].cells[1].text = access
        table.rows[i].cells[2].text = modules

    # 11. API Module Reference
    add_heading(doc, "11. API Module Reference", 1)
    api_table = doc.add_table(rows=13, cols=2)
    api_table.style = "Table Grid"
    api_table.rows[0].cells[0].text = "Module"
    api_table.rows[0].cells[1].text = "Endpoints"
    api_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    api_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    apis = [
        ("Auth", "POST /api/auth/login, GET /api/auth/me"),
        ("Bookings", "/api/bookings — CRUD"),
        ("LR", "/api/lr — CRUD"),
        ("Vehicles", "/api/vehicles — CRUD"),
        ("Drivers", "/api/drivers — CRUD"),
        ("Customers", "/api/customers — CRUD"),
        ("Vendors", "/api/vendors — CRUD"),
        ("Expenses", "/api/expenses — CRUD"),
        ("Dashboard", "/api/dashboard/stats, charts, alerts"),
        ("Reports", "/api/reports/trips, income, expenses, cash-flow"),
        ("Accounting", "/api/accounting/* — registers, statements, vouchers"),
        ("Settings", "GET/PUT /api/settings"),
    ]
    for i, (mod, ep) in enumerate(apis, 1):
        api_table.rows[i].cells[0].text = mod
        api_table.rows[i].cells[1].text = ep

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer.add_run("— End of Document —")
    frun.italic = True
    frun.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")

if __name__ == "__main__":
    main()
